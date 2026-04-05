import streamlit as st
from docx import Document
import difflib
import re
import unicodedata

# ─────────────────────────────────────────────
# FILE READING
# ─────────────────────────────────────────────

def get_text_file(file):
    content = file.read().decode('utf-8')
    lines = content.split('\n')
    return [line.strip() for line in lines if line.strip()]

# ─────────────────────────────────────────────
# TEXT NORMALIZATION
# ─────────────────────────────────────────────

def flatten_poem_lines(paragraphs, window_size=4):
    result = []
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        if len(para) < 50:
            poem_lines = [para]
            j = i + 1
            while j < len(paragraphs) and len(paragraphs[j]) < 50 and j < i + 20:
                poem_lines.append(paragraphs[j])
                j += 1
            if len(poem_lines) >= window_size:
                result.append(' '.join(poem_lines))
                i = j
            else:
                result.append(para)
                i += 1
        else:
            result.append(para)
            i += 1
    return result

def simplify(text):
    """Normalize for comparison only: remove punctuation, hyphens, lowercase."""
    if not text:
        return ""
    text = unicodedata.normalize('NFC', text)
    text = re.sub(r'\(\d+\)', '', text)
    text = text.replace('-', ' ')
    text = re.sub(r'[^\w\s]', '', text, flags=re.UNICODE)
    text = text.lower()
    # Map visually-similar character variants to canonical form
    text = text.replace('đ', 'd')  # Latin Small Letter D with stroke -> d
    text = text.replace('ð', 'd')  # Latin Small Letter Eth -> d
    return ' '.join(text.split())

# ─────────────────────────────────────────────
# SENTENCE SPLITTING
# The key insight: diff at SENTENCE level, not paragraph level.
# Paragraphs are just display containers — sentences are the real units of meaning.
# SequenceMatcher works reliably when each token is a sentence,
# because sentences are long enough to be uniquely identifiable,
# and short enough that boundary shifts don't hide matches.
# ─────────────────────────────────────────────

def split_into_sentences(paragraphs):
    """
    Split paragraphs into individual sentences.
    Returns list of (sentence_text, para_idx) so we can restore
    paragraph breaks in the display.
    """
    result = []
    for para_idx, para in enumerate(paragraphs):
        # Split at sentence-ending punctuation followed by whitespace
        parts = re.split(r'(?<=[.;!?])\s+', para.strip())
        for part in parts:
            part = part.strip()
            if part:
                result.append((part, para_idx))
    return result


def pair_sentences(raw_paragraphs, proc_paragraphs):
    """Return aligned sentence triples for display and comparison."""
    raw_sentences = split_into_sentences(raw_paragraphs)
    proc_sentences = split_into_sentences(proc_paragraphs)

    paired = []
    for idx, (raw_text, para_idx) in enumerate(raw_sentences):
        proc_text = proc_sentences[idx][0] if idx < len(proc_sentences) else raw_text
        paired.append((raw_text, simplify(proc_text), para_idx))

    for idx in range(len(raw_sentences), len(proc_sentences)):
        proc_text, para_idx = proc_sentences[idx]
        paired.append((proc_text, simplify(proc_text), para_idx))

    return paired

# ─────────────────────────────────────────────
# DIFF AT SENTENCE LEVEL
# ─────────────────────────────────────────────

def build_diffs(sentences_a, sentences_b):
    """
    Diff two sentence lists. Each sentence is (raw_text, normalized_compare_text, para_idx).
    Uses SequenceMatcher on the normalized compare_text values at sentence level.
    Returns list of diff blocks for non-equal regions.
    """
    keys_a = [compare_text for _, compare_text, _ in sentences_a]
    keys_b = [compare_text for _, compare_text, _ in sentences_b]

    matcher = difflib.SequenceMatcher(None, keys_a, keys_b, autojunk=False)
    diffs = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
        block_a = sentences_a[i1:i2]
        block_b = sentences_b[j1:j2]
        
        # Skip blocks where the joined normalized content is identical
        # (e.g., formatting/boundary differences but same semantic content)
        joined_a = ' '.join(compare_text for _, compare_text, _ in block_a)
        joined_b = ' '.join(compare_text for _, compare_text, _ in block_b)
        
        if joined_a == joined_b:
            continue
        
        diffs.append({
            'sents_a': block_a,
            'sents_b': block_b,
            'start_a': i1,
            'start_b': j1,
            'end_a': i2,
            'end_b': j2,
        })

    return diffs

# ─────────────────────────────────────────────
# WORD-LEVEL HIGHLIGHT within a sentence pair
# ─────────────────────────────────────────────

def highlight_words(text_a, text_b):
    """Word-level diff highlight between two text blocks.
    Preserves original text (with hyphens and punctuation) for display."""
    if simplify(text_a) == simplify(text_b):
        return text_a, text_b

    # Replace newlines with a placeholder token to preserve them
    NEWLINE_MARKER = '__PARA_BREAK__'
    text_a_marked = text_a.replace('\n', ' ' + NEWLINE_MARKER + ' ')
    text_b_marked = text_b.replace('\n', ' ' + NEWLINE_MARKER + ' ')
    
    # Work with original tokens (preserves hyphens in multi-part words)
    words_a_original = text_a_marked.split()
    words_b_original = text_b_marked.split()
    
    # Create hyphen-replaced versions for normalization/comparison
    words_a_for_norm = [w.replace('-', ' ') for w in words_a_original]
    words_b_for_norm = [w.replace('-', ' ') for w in words_b_original]
    
    # Normalize and track changed indices at word level
    norm_a = ' '.join(simplify(w) for w in words_a_for_norm)
    norm_b = ' '.join(simplify(w) for w in words_b_for_norm)
    
    if norm_a == norm_b:
        return text_a, text_b
    
    norm_words_a = norm_a.split()
    norm_words_b = norm_b.split()
    
    matcher = difflib.SequenceMatcher(None, norm_words_a, norm_words_b, autojunk=False)
    
    # Track which normalized words are changed
    changed_word_indices_a = set()
    changed_word_indices_b = set()
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag != 'equal':
            for idx in range(i1, i2):
                changed_word_indices_a.add(idx)
            for idx in range(j1, j2):
                changed_word_indices_b.add(idx)
    
    # Apply highlighting to original tokens (with hyphens preserved)
    def highlight_original_tokens(original_words, norm_for_comparison, changed_indices, is_b_side=False):
        """Apply highlights to original tokens based on normalized comparison indices."""
        norm_idx = 0
        result = []
        
        color_a = "background-color:#ffcccc;"
        color_b = "background-color:#ccffcc; color:black; font-weight:bold;"
        color = color_b if is_b_side else color_a
        
        for orig_word in original_words:
            if orig_word == NEWLINE_MARKER:
                result.append(orig_word)
                norm_idx += 1  # Account for marker in normalized list
                continue
            
            # Normalize the original word (replace hyphens with spaces and simplify)
            norm_version = simplify(orig_word.replace('-', ' '))
            if not norm_version:
                result.append(orig_word)
                continue
            
            # Count how many tokens this word represents when normalized
            num_tokens = len(norm_version.split())
            
            # Check if any of these tokens are in the changed set
            is_changed = any(norm_idx + i in changed_indices for i in range(num_tokens))
            
            if is_changed:
                result.append(f"<span style='{color}'>{orig_word}</span>")
            else:
                result.append(orig_word)
            
            norm_idx += num_tokens
        
        return ' '.join(result)
    
    out_a = highlight_original_tokens(words_a_original, words_a_for_norm, changed_word_indices_a, is_b_side=False)
    out_b = highlight_original_tokens(words_b_original, words_b_for_norm, changed_word_indices_b, is_b_side=True)
    
    # Restore newlines from markers
    out_a = out_a.replace(' ' + NEWLINE_MARKER + ' ', '\n')
    out_b = out_b.replace(' ' + NEWLINE_MARKER + ' ', '\n')
    
    return out_a, out_b

# ─────────────────────────────────────────────
# RENDERING
# ─────────────────────────────────────────────

def render_sentence_list(sentences, highlight_map=None):
    """
    Render a list of sentences as HTML, inserting paragraph breaks on para_idx changes.
    highlight_map: dict of sentence_idx -> highlighted_html (overrides plain text).
    """
    html = []
    last_para = None
    for idx, item in enumerate(sentences):
        if len(item) == 3:
            text, _, para_idx = item
        else:
            text, para_idx = item
        if last_para is not None and para_idx != last_para:
            html.append('<br><br>')
        last_para = para_idx
        if highlight_map and idx in highlight_map:
            html.append(highlight_map[idx] + ' ')
        else:
            html.append(text + ' ')
    return ''.join(html)

def render_diff(diff, sentences_a, sentences_b, context_lines):
    """
    Render one diff block with word-level highlights.
    Join all sentences in the block, restore line breaks, then compare and highlight.
    """
    sents_a = diff['sents_a']
    sents_b = diff['sents_b']

    # Reconstruct sentences with line breaks restored
    def reconstruct_with_breaks(sentences):
        """Reconstruct text with line breaks restored at para_idx boundaries."""
        parts = []
        last_para = None
        for raw_text, _, para_idx in sentences:
            if last_para is not None and para_idx != last_para:
                parts.append('\n')
            last_para = para_idx
            parts.append(raw_text)
        return ' '.join(parts)
    
    # Reconstruct both sides with line breaks
    text_a = reconstruct_with_breaks(sents_a)
    text_b = reconstruct_with_breaks(sents_b)
    
    # Get normalized versions for comparison
    joined_norm_a = ' '.join(compare_text for _, compare_text, _ in sents_a)
    joined_norm_b = ' '.join(compare_text for _, compare_text, _ in sents_b)
    
    # If normalized versions are identical, show without highlighting
    if joined_norm_a == joined_norm_b:
        high_a = text_a.replace('\n', '<br><br>')
        high_b = text_b.replace('\n', '<br><br>')
    else:
        # Do word-level diff on the full text (preserving line breaks as markers)
        high_a, high_b = highlight_words(text_a, text_b)
        # Replace line break markers with HTML breaks
        high_a = high_a.replace('\n', '<br><br>')
        high_b = high_b.replace('\n', '<br><br>')

    # Context sentences after this diff block
    end_a, end_b = diff['end_a'], diff['end_b']
    ctx_a = sentences_a[end_a:end_a + context_lines]
    ctx_b = sentences_b[end_b:end_b + context_lines]

    if ctx_a or ctx_b:
        sep = "<br><br><div style='border-top:1px dashed #ccc; margin:15px 0; padding-top:15px; color:#999;'>"
        high_a += sep + render_sentence_list(ctx_a) + "</div>"
        high_b += sep + render_sentence_list(ctx_b) + "</div>"

    return high_a, high_b

# ─────────────────────────────────────────────
# APP
# ─────────────────────────────────────────────

st.set_page_config(page_title="Deep-Sync Comparison", layout="wide")
st.title("Document Comparison (Typo & Offset Resilient)")

with st.sidebar:
    st.markdown("---")
    st.markdown("<small style='color:#888;'>**Version:** 2025-04-03 (Fixed highlight algorithm)</small>", unsafe_allow_html=True)
    st.markdown("---")
    f_orig = st.file_uploader("Original Document", type=["txt", "docx"])
    f_rev  = st.file_uploader("Revised Document",  type=["txt", "docx"])
    anchor = st.text_input("Anchor Point", "Như vậy tôi nghe")
    context_lines = st.number_input("Context sentences", min_value=0, max_value=50, value=5)
    debug_mode = st.checkbox("Debug mode (show normalized text)", value=False)

if f_orig and f_rev:
    # ── Read ──
    if f_orig.name.endswith('.txt'):
        raw_a = get_text_file(f_orig)
    else:
        raw_a = [p.text.strip() for p in Document(f_orig).paragraphs if p.text.strip()]

    if f_rev.name.endswith('.txt'):
        raw_b = get_text_file(f_rev)
    else:
        raw_b = [p.text.strip() for p in Document(f_rev).paragraphs if p.text.strip()]

    # ── Flatten poem lines FIRST (before any processing) ──
    raw_a = flatten_poem_lines(raw_a)
    raw_b = flatten_poem_lines(raw_b)
    
    # ── THEN apply hyphen normalization for comparison ──
    proc_a = [p.replace('-', ' ') for p in raw_a]
    proc_b = [p.replace('-', ' ') for p in raw_b]

    # ── Anchor alignment (use proc for normalization, raw for output) ──
    idx_a = next((i for i, x in enumerate(proc_a) if simplify(anchor) in simplify(x)), 0)
    idx_b = next((i for i, x in enumerate(proc_b) if simplify(anchor) in simplify(x)), 0)
    text_a = raw_a[idx_a:]
    text_b = raw_b[idx_b:]
    # ── Compute full sentences and diffs for global navigation ──
    full_sentences_a = pair_sentences(raw_a, proc_a)
    full_sentences_b = pair_sentences(raw_b, proc_b)
    full_diffs = build_diffs(full_sentences_a, full_sentences_b)
    pre_sentences_a = split_into_sentences(raw_a[:idx_a])
    pre_sentences_b = split_into_sentences(raw_b[:idx_b])
    start_a = len(pre_sentences_a)
    start_b = len(pre_sentences_b)

    # Find initial global diff index: prefer diff containing anchor, else first after anchor
    start_diff_idx = None
    for i, d in enumerate(full_diffs):
        for raw_text, _, _ in d['sents_a'] + d['sents_b']:
            if simplify(anchor) in simplify(raw_text):
                start_diff_idx = i
                break
        if start_diff_idx is not None:
            break
    if start_diff_idx is None:
        start_diff_idx = next(
            (i for i, d in enumerate(full_diffs)
             if d['end_a'] > start_a or d['end_b'] > start_b),
            0
        )

    # ── Split to sentences after anchor for display context range if needed ──
    sentences_a = pair_sentences(text_a, proc_a[idx_a:])
    sentences_b = pair_sentences(text_b, proc_b[idx_b:])

    # ── Session state ──
    anchor_changed = "last_anchor" in st.session_state and st.session_state.last_anchor != anchor
    if "last_anchor" not in st.session_state:
        st.session_state.last_anchor = anchor
    elif anchor_changed:
        st.session_state.last_anchor = anchor
        st.session_state.nav = start_diff_idx
        st.session_state.jump = str(start_diff_idx + 1)

    if "nav" not in st.session_state:
        st.session_state.nav = start_diff_idx
    if "jump" not in st.session_state:
        st.session_state.jump = str(st.session_state.nav + 1)

    # callback for jump field changes
    def on_jump_change():
        try:
            jump_val = int(st.session_state.jump)
            if 1 <= jump_val <= len(full_diffs):
                st.session_state.nav = jump_val - 1
        except ValueError:
            pass

    # ── Use full diffs for global navigation and rendering ──
    if not full_diffs:
        st.success("Documents are fully synchronized. No content differences found!")
    else:
        if "nav" not in st.session_state:
            st.session_state.nav = start_diff_idx

        # Clamp stale nav index into bounds
        st.session_state.nav = min(max(st.session_state.nav, 0), len(full_diffs) - 1)
        st.session_state.jump = str(st.session_state.nav + 1)

        # Navigation
        c1, c2, c3 = st.columns([1, 2, 1])
        with c1:
            if st.button("Previous"):
                st.session_state.nav = (st.session_state.nav - 1) % len(full_diffs)

        with c2:
            st.markdown(
                f"<p style='text-align:center'>Difference "
                f"<b>{st.session_state.nav + 1}</b> of <b>{len(full_diffs)}</b></p>",
                unsafe_allow_html=True
            )
            st.text_input("Jump to difference", key="jump", on_change=on_jump_change)

        with c3:
            if st.button("Next"):
                st.session_state.nav = (st.session_state.nav + 1) % len(full_diffs)

        high_a, high_b = render_diff(
            full_diffs[st.session_state.nav], full_sentences_a, full_sentences_b, context_lines
        )

        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Original")
            st.markdown(
                f"<div style='border:1px solid #ddd; padding:20px; "
                f"font-size:18px; min-height:300px; line-height:1.9;'>{high_a}</div>",
                unsafe_allow_html=True
            )
            if debug_mode:
                st.write("**Sentences in diff block A (normalized):**")
                for i, (raw, norm, pi) in enumerate(full_diffs[st.session_state.nav]['sents_a']):
                    st.caption(f"[{i}] Para {pi}: {norm}")
        with col2:
            st.subheader("Revised")
            st.markdown(
                f"<div style='border:1px solid #ddd; padding:20px; "
                f"font-size:18px; min-height:300px; line-height:1.9;'>{high_b}</div>",
                unsafe_allow_html=True
            )
            if debug_mode:
                st.write("**Sentences in diff block B (normalized):**")
                for i, (raw, norm, pi) in enumerate(full_diffs[st.session_state.nav]['sents_b']):
                    st.caption(f"[{i}] Para {pi}: {norm}")